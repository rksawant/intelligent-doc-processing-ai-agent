from fpdf import FPDF  # provided by fpdf2
from docx import Document
import os

os.makedirs("test_data", exist_ok=True)

# Create DOCX sample
doc = Document()
doc.add_heading("AI Document Processing Test", 0)
doc.add_paragraph("This document is used to verify the LangChain + Pinecone pipeline.")
doc.save("test_data/sample_doc.docx")

# Create PDF sample
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, txt="AI Agent PDF Test Document", ln=True, align="C")
pdf.output("test_data/sample_pdf.pdf")

print("✅ Test data created in ./test_data/")
