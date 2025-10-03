"""
Document Processing Agent for handling various document formats
"""
import os
import logging
from typing import List, Dict, Any, Optional, BinaryIO
from datetime import datetime
import hashlib
import mimetypes

# Document processing libraries
import PyPDF2
from docx import Document
import html2text

# AWS services
from services import S3Service, TextractService, BedrockService
from config.config import document_config, aws_config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Document processing agent for various file formats"""
    
    def __init__(self):
        self.s3_service = S3Service()
        self.textract_service = TextractService()
        self.bedrock_service = BedrockService()
        self.supported_formats = document_config.supported_formats
        self.max_file_size = document_config.max_file_size_mb * 1024 * 1024  # Convert to bytes
    
    def process_document(self, file_path: str, metadata: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        Process a document and extract text content
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata to attach
            
        Returns:
            Processing result with extracted content and metadata
        """
        try:
            # Validate file
            if not self._validate_file(file_path):
                return {'error': 'Invalid file format or size'}
            
            # Upload to S3
            s3_key = self._upload_to_s3(file_path, metadata)
            if not s3_key:
                return {'error': 'Failed to upload document to S3'}
            
            # Extract text based on file type
            file_extension = os.path.splitext(file_path)[1].lower()
            extraction_result = self._extract_text_by_type(file_path, s3_key, file_extension)
            
            if 'error' in extraction_result:
                return extraction_result
            
            # Generate document ID and metadata
            document_id = self._generate_document_id(file_path, extraction_result['text'])
            processing_metadata = {
                'document_id': document_id,
                'original_filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'processing_timestamp': datetime.utcnow().isoformat(),
                'file_type': file_extension,
                **(metadata or {})
            }
            
            # Upload processed content to S3
            processed_s3_key = self.s3_service.upload_processed_document(
                extraction_result['text'],
                os.path.basename(file_path),
                processing_metadata
            )
            
            return {
                'success': True,
                'document_id': document_id,
                's3_key': s3_key,
                'processed_s3_key': processed_s3_key,
                'text_content': extraction_result['text'],
                'metadata': processing_metadata,
                'structured_data': extraction_result.get('structured_data', {}),
                'word_count': len(extraction_result['text'].split()),
                'character_count': len(extraction_result['text'])
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {'error': str(e)}
    
    def process_document_from_bytes(self, file_bytes: bytes, filename: str, 
                                  metadata: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        Process document from bytes
        
        Args:
            file_bytes: Document bytes
            filename: Original filename
            metadata: Optional metadata
            
        Returns:
            Processing result
        """
        try:
            # Validate file size
            if len(file_bytes) > self.max_file_size:
                return {'error': f'File size exceeds maximum limit of {self.max_file_size} bytes'}
            
            # Determine file type
            file_extension = os.path.splitext(filename)[1].lower()
            if not self._is_supported_format(file_extension):
                return {'error': f'Unsupported file format: {file_extension}'}
            
            # Upload to S3
            s3_key = self._upload_bytes_to_s3(file_bytes, filename, metadata)
            if not s3_key:
                return {'error': 'Failed to upload document to S3'}
            
            # Extract text
            extraction_result = self._extract_text_from_bytes(file_bytes, file_extension)
            
            if 'error' in extraction_result:
                return extraction_result
            
            # Generate document ID and metadata
            document_id = self._generate_document_id(filename, extraction_result['text'])
            processing_metadata = {
                'document_id': document_id,
                'original_filename': filename,
                'file_size': len(file_bytes),
                'processing_timestamp': datetime.utcnow().isoformat(),
                'file_type': file_extension,
                **(metadata or {})
            }
            
            # Upload processed content
            processed_s3_key = self.s3_service.upload_processed_document(
                extraction_result['text'],
                filename,
                processing_metadata
            )
            
            return {
                'success': True,
                'document_id': document_id,
                's3_key': s3_key,
                'processed_s3_key': processed_s3_key,
                'text_content': extraction_result['text'],
                'metadata': processing_metadata,
                'structured_data': extraction_result.get('structured_data', {}),
                'word_count': len(extraction_result['text'].split()),
                'character_count': len(extraction_result['text'])
            }
            
        except Exception as e:
            logger.error(f"Error processing document from bytes: {e}")
            return {'error': str(e)}
    
    def _validate_file(self, file_path: str) -> bool:
        """Validate file format and size"""
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                logger.error(f"File size {file_size} exceeds maximum limit {self.max_file_size}")
                return False
            
            # Check file format
            file_extension = os.path.splitext(file_path)[1].lower()
            if not self._is_supported_format(file_extension):
                logger.error(f"Unsupported file format: {file_extension}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {e}")
            return False
    
    def _is_supported_format(self, file_extension: str) -> bool:
        """Check if file format is supported"""
        return file_extension.lstrip('.') in self.supported_formats
    
    def _upload_to_s3(self, file_path: str, metadata: Optional[Dict[str, str]] = None) -> Optional[str]:
        """Upload file to S3"""
        try:
            filename = os.path.basename(file_path)
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            s3_key = f"{aws_config.s3_raw_prefix}{timestamp}_{filename}"
            
            success = self.s3_service.upload_document(file_path, s3_key, metadata)
            return s3_key if success else None
            
        except Exception as e:
            logger.error(f"Error uploading to S3: {e}")
            return None
    
    def _upload_bytes_to_s3(self, file_bytes: bytes, filename: str, 
                           metadata: Optional[Dict[str, str]] = None) -> Optional[str]:
        """Upload file bytes to S3"""
        try:
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            s3_key = f"{aws_config.s3_raw_prefix}{timestamp}_{filename}"
            
            from io import BytesIO
            file_obj = BytesIO(file_bytes)
            
            success = self.s3_service.upload_file_object(file_obj, s3_key, metadata)
            return s3_key if success else None
            
        except Exception as e:
            logger.error(f"Error uploading bytes to S3: {e}")
            return None
    
    def _extract_text_by_type(self, file_path: str, s3_key: str, file_extension: str) -> Dict[str, Any]:
        """Extract text based on file type"""
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path, s3_key)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.html':
                return self._extract_from_html(file_path)
            else:
                return {'error': f'No extraction method for file type: {file_extension}'}
                
        except Exception as e:
            logger.error(f"Error extracting text by type: {e}")
            return {'error': str(e)}
    
    def _extract_text_from_bytes(self, file_bytes: bytes, file_extension: str) -> Dict[str, Any]:
        """Extract text from file bytes"""
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf_bytes(file_bytes)
            elif file_extension == '.docx':
                return self._extract_from_docx_bytes(file_bytes)
            elif file_extension == '.txt':
                return self._extract_from_txt_bytes(file_bytes)
            elif file_extension == '.html':
                return self._extract_from_html_bytes(file_bytes)
            else:
                return {'error': f'No extraction method for file type: {file_extension}'}
                
        except Exception as e:
            logger.error(f"Error extracting text from bytes: {e}")
            return {'error': str(e)}
    
    def _extract_from_pdf(self, file_path: str, s3_key: str) -> Dict[str, Any]:
        """Extract text from PDF using Textract"""
        try:
            # Try Textract first for better OCR capabilities
            textract_result = self.textract_service.extract_structured_data(s3_key)
            
            if 'error' not in textract_result:
                return {
                    'text': textract_result['text'],
                    'structured_data': {
                        'tables': textract_result.get('tables', []),
                        'forms': textract_result.get('forms', []),
                        'queries': textract_result.get('queries', [])
                    }
                }
            else:
                # Fallback to PyPDF2
                return self._extract_from_pdf_pypdf2(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting from PDF with Textract: {e}")
            # Fallback to PyPDF2
            return self._extract_from_pdf_pypdf2(file_path)
    
    def _extract_from_pdf_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                
                return {'text': text.strip()}
                
        except Exception as e:
            logger.error(f"Error extracting from PDF with PyPDF2: {e}")
            return {'error': str(e)}
    
    def _extract_from_pdf_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from PDF bytes using PyPDF2"""
        try:
            from io import BytesIO
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return {'text': text.strip()}
            
        except Exception as e:
            logger.error(f"Error extracting from PDF bytes: {e}")
            return {'error': str(e)}
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {'text': text.strip()}
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            return {'error': str(e)}
    
    def _extract_from_docx_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from DOCX bytes"""
        try:
            from io import BytesIO
            doc_file = BytesIO(file_bytes)
            doc = Document(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {'text': text.strip()}
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX bytes: {e}")
            return {'error': str(e)}
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {'text': text}
            
        except Exception as e:
            logger.error(f"Error extracting from TXT: {e}")
            return {'error': str(e)}
    
    def _extract_from_txt_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from TXT bytes"""
        try:
            text = file_bytes.decode('utf-8')
            return {'text': text}
            
        except Exception as e:
            logger.error(f"Error extracting from TXT bytes: {e}")
            return {'error': str(e)}
    
    def _extract_from_html(self, file_path: str) -> Dict[str, Any]:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            return self._extract_from_html_content(html_content)
            
        except Exception as e:
            logger.error(f"Error extracting from HTML: {e}")
            return {'error': str(e)}
    
    def _extract_from_html_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from HTML bytes"""
        try:
            html_content = file_bytes.decode('utf-8')
            return self._extract_from_html_content(html_content)
            
        except Exception as e:
            logger.error(f"Error extracting from HTML bytes: {e}")
            return {'error': str(e)}
    
    def _extract_from_html_content(self, html_content: str) -> Dict[str, Any]:
        """Extract text from HTML content"""
        try:
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = True
            h.ignore_emphasis = False
            
            text = h.handle(html_content)
            return {'text': text.strip()}
            
        except Exception as e:
            logger.error(f"Error extracting from HTML content: {e}")
            return {'error': str(e)}
    
    def _generate_document_id(self, filename: str, content: str) -> str:
        """Generate unique document ID"""
        try:
            # Create hash from filename and content
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            filename_hash = hashlib.md5(filename.encode('utf-8')).hexdigest()
            
            # Combine hashes
            combined_hash = hashlib.md5(f"{filename_hash}{content_hash}".encode('utf-8')).hexdigest()
            
            # Add timestamp for uniqueness
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            
            return f"doc_{timestamp}_{combined_hash[:12]}"
            
        except Exception as e:
            logger.error(f"Error generating document ID: {e}")
            # Fallback to timestamp-based ID
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S_%f")
            return f"doc_{timestamp}"
    
    def get_document_summary(self, document_id: str, max_length: int = 500) -> Dict[str, Any]:
        """
        Get document summary using Bedrock
        
        Args:
            document_id: Document ID
            max_length: Maximum summary length
            
        Returns:
            Document summary
        """
        try:
            # Get document content from S3
            s3_key = f"{aws_config.s3_processed_prefix}{document_id}.txt"
            content = self.s3_service.get_document_content(s3_key)
            
            if not content:
                return {'error': 'Document not found'}
            
            text_content = content.decode('utf-8')
            
            # Generate summary using Bedrock
            summary = self.bedrock_service.summarize_document(text_content, max_length)
            
            return {
                'document_id': document_id,
                'summary': summary,
                'original_length': len(text_content),
                'summary_length': len(summary)
            }
            
        except Exception as e:
            logger.error(f"Error getting document summary: {e}")
            return {'error': str(e)}
    
    def extract_key_information(self, document_id: str, information_types: List[str]) -> Dict[str, Any]:
        """
        Extract key information from document
        
        Args:
            document_id: Document ID
            information_types: List of information types to extract
            
        Returns:
            Extracted information
        """
        try:
            # Get document content
            s3_key = f"{aws_config.s3_processed_prefix}{document_id}.txt"
            content = self.s3_service.get_document_content(s3_key)
            
            if not content:
                return {'error': 'Document not found'}
            
            text_content = content.decode('utf-8')
            
            # Extract information using Bedrock
            extracted_info = self.bedrock_service.extract_key_information(text_content, information_types)
            
            return {
                'document_id': document_id,
                'extracted_information': extracted_info,
                'information_types': information_types
            }
            
        except Exception as e:
            logger.error(f"Error extracting key information: {e}")
            return {'error': str(e)}
